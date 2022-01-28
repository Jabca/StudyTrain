import os
from math import ceil

import docx
from docx.shared import Inches, Cm


def save_dir_to_docx(folder_path, save_path):
    document = docx.Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)

    images_names = list(os.listdir(folder_path))

    tbl = document.add_table(rows=ceil(len(images_names) / 2) - 1, cols=2)
    row_cells = tbl.add_row().cells
    for i, name in enumerate(images_names):
        paragraph = row_cells[i % 2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(f"{folder_path}/{name}", width=Inches(4))
    document.save(save_path)

